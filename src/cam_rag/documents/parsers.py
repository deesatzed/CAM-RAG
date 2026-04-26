"""Document parser backends with optional dependencies."""

from __future__ import annotations

import re
from pathlib import Path


class MissingParserDependencyError(RuntimeError):
    """Raised when a parser extra is needed but not installed."""


def parse_pdf_text(path: str | Path) -> tuple[str, dict]:
    """Extract markdown-ish text from a PDF.

    Uses `pymupdf4llm` when available because it preserves useful headings for
    RAG chunking. Falls back to PyMuPDF plain text if only `pymupdf` is present.
    """

    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"PDF not found: {file_path}")

    try:
        import pymupdf4llm  # type: ignore[import-not-found]
    except ImportError:
        pymupdf4llm = None

    if pymupdf4llm is not None:
        text = pymupdf4llm.to_markdown(str(file_path))
        page_count = _pdf_page_count(file_path)
        return text, {"page_count": page_count, "parser": "pymupdf4llm"}

    try:
        import pymupdf  # type: ignore[import-not-found]
    except ImportError as exc:
        raise MissingParserDependencyError(
            "PDF ingestion requires the `pdf` extra: install pymupdf4llm or pymupdf."
        ) from exc

    doc = pymupdf.open(str(file_path))
    try:
        text = "\n\n".join(page.get_text("text") for page in doc)
        page_count = len(doc)
    finally:
        doc.close()
    return text, {"page_count": page_count, "parser": "pymupdf"}


def parse_docx_text(path: str | Path) -> tuple[str, dict]:
    """Extract text from a DOCX file using `python-docx`."""

    file_path = Path(path)
    if not file_path.exists():
        raise FileNotFoundError(f"DOCX not found: {file_path}")

    try:
        from docx import Document as DocxDocument  # type: ignore[import-not-found]
    except ImportError as exc:
        raise MissingParserDependencyError(
            "DOCX ingestion requires the `docx` extra: install python-docx."
        ) from exc

    doc = DocxDocument(str(file_path))
    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n\n".join(parts), {"page_count": 0, "parser": "python-docx"}


def extract_title(text: str, fallback: str) -> str:
    """Extract a useful title from markdown-ish text."""

    for line in text.strip().splitlines()[:20]:
        stripped = line.strip()
        if stripped.startswith("#"):
            title = stripped.lstrip("#").strip()
            if title:
                return title[:120]
    for line in text.strip().splitlines()[:10]:
        cleaned = line.strip().strip("#").strip()
        if cleaned and len(cleaned) > 3:
            return re.sub(r"\s+", " ", cleaned)[:120]
    return fallback


def _pdf_page_count(path: Path) -> int:
    try:
        import pymupdf  # type: ignore[import-not-found]
    except ImportError:
        return 0
    doc = pymupdf.open(str(path))
    try:
        return len(doc)
    finally:
        doc.close()
