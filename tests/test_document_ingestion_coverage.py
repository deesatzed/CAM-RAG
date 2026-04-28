"""Close coverage gaps for parsers.py, chunking.py, and folder.py.

Targets
-------
- parsers.py  54% -> 90%+
- chunking.py 64% -> 90%+
- folder.py   78% -> 90%+

Every test uses real files written to ``tmp_path``; no mocks.
"""

from __future__ import annotations

import builtins
import json
import logging
import os
import sys
from pathlib import Path

import pytest

from cam_rag.documents.chunking import (
    _split_long_text,
    chunk_document,
    chunk_documents,
)
from cam_rag.documents.folder import (
    _first_string,
    _read_json_documents,
    _read_parsed_document,
    _safe_read_parsed,
    _safe_read_text,
    read_document_folder,
)
from cam_rag.documents.parsers import (
    _pdf_page_count,
    extract_title,
    parse_docx_text,
    parse_pdf_text,
)
from cam_rag.rag.models import CorpusDocument
from cam_rag.rag.spec import RAGAppSpec

# ======================================================================
# helpers
# ======================================================================


def _make_doc(text: str, doc_id: str = "doc_test") -> CorpusDocument:
    return CorpusDocument(
        id=doc_id, text=text, source="test.md", title="Test"
    )


def _create_minimal_pdf(file_path: Path, text: str = "Hello PDF World") -> Path:
    """Create a real minimal PDF using pymupdf."""
    import pymupdf

    doc = pymupdf.open()
    page = doc.new_page()
    page.insert_text((72, 72), text)
    doc.save(str(file_path))
    doc.close()
    return file_path


def _create_minimal_docx(file_path: Path, paragraphs: list[str] | None = None) -> Path:
    """Create a real minimal DOCX using python-docx."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    for text in (paragraphs or ["Hello DOCX World"]):
        doc.add_paragraph(text)
    doc.save(str(file_path))
    return file_path


def _create_docx_with_table(
    file_path: Path,
    paragraphs: list[str],
    table_rows: list[list[str]],
) -> Path:
    """Create a real DOCX containing both paragraphs and a table."""
    from docx import Document as DocxDocument

    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        num_cols = len(table_rows[0])
        table = doc.add_table(rows=len(table_rows), cols=num_cols)
        for ri, row_data in enumerate(table_rows):
            for ci, cell_text in enumerate(row_data):
                table.rows[ri].cells[ci].text = cell_text
    doc.save(str(file_path))
    return file_path


# ======================================================================
# parsers.py -- PDF happy paths (lines 22, 30-32, 41-47, 97-105)
# ======================================================================


class TestParsePdfText:
    """Tests for ``parse_pdf_text`` using real pymupdf / pymupdf4llm."""

    def test_file_not_found_raises(self, tmp_path: Path):
        """Line 22: FileNotFoundError when the PDF does not exist."""
        with pytest.raises(FileNotFoundError, match="PDF not found"):
            parse_pdf_text(tmp_path / "nonexistent.pdf")

    def test_happy_path_with_pymupdf4llm(self, tmp_path: Path):
        """Lines 30-32: pymupdf4llm branch returns markdown text + metadata."""
        pdf = _create_minimal_pdf(tmp_path / "sample.pdf", text="RAG test content")
        text, meta = parse_pdf_text(pdf)
        assert "RAG test content" in text
        assert meta["parser"] in {"pymupdf4llm", "pymupdf"}
        assert isinstance(meta["page_count"], int)
        assert meta["page_count"] >= 1

    def test_happy_path_returns_nonempty_text(self, tmp_path: Path):
        """The parsed text is a non-empty string."""
        pdf = _create_minimal_pdf(tmp_path / "nonempty.pdf", text="Content here")
        text, _ = parse_pdf_text(pdf)
        assert len(text.strip()) > 0

    def test_pymupdf_fallback_when_pymupdf4llm_absent(
        self, monkeypatch, tmp_path: Path
    ):
        """Lines 41-47: pymupdf fallback when pymupdf4llm is not installed.

        Uses the same import-interception technique as the existing
        test_document_parsers.py to simulate the optional dependency
        being absent while keeping pymupdf itself available.
        """
        pdf = _create_minimal_pdf(tmp_path / "fallback.pdf", text="Fallback content")
        real_import = builtins.__import__

        def _block_pymupdf4llm(name, *args, **kwargs):
            if name == "pymupdf4llm":
                raise ImportError(name)
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", _block_pymupdf4llm)
        text, meta = parse_pdf_text(pdf)
        assert "Fallback content" in text
        assert meta["parser"] == "pymupdf"
        assert meta["page_count"] >= 1


class TestPdfPageCount:
    """Tests for ``_pdf_page_count`` (lines 97-105)."""

    def test_returns_page_count_for_real_pdf(self, tmp_path: Path):
        pdf = _create_minimal_pdf(tmp_path / "pages.pdf")
        count = _pdf_page_count(pdf)
        assert count == 1

    def test_multipage_pdf(self, tmp_path: Path):
        """A PDF with two pages returns 2."""
        import pymupdf

        doc = pymupdf.open()
        doc.new_page()
        doc.new_page()
        path = tmp_path / "two_pages.pdf"
        doc.save(str(path))
        doc.close()

        assert _pdf_page_count(path) == 2

    def test_returns_zero_when_pymupdf_absent(self, monkeypatch, tmp_path: Path):
        """Lines 99-100: returns 0 when pymupdf is not importable."""
        pdf = _create_minimal_pdf(tmp_path / "count.pdf")
        real_import = builtins.__import__

        def _block_pymupdf(name, *args, **kwargs):
            if name == "pymupdf":
                raise ImportError(name)
            return real_import(name, *args, **kwargs)

        monkeypatch.setattr(builtins, "__import__", _block_pymupdf)
        assert _pdf_page_count(pdf) == 0


# ======================================================================
# parsers.py -- DOCX happy paths (lines 55, 64-77)
# ======================================================================


class TestParseDocxText:
    """Tests for ``parse_docx_text`` using real python-docx files."""

    def test_file_not_found_raises(self, tmp_path: Path):
        """Line 55: FileNotFoundError when the DOCX does not exist."""
        with pytest.raises(FileNotFoundError, match="DOCX not found"):
            parse_docx_text(tmp_path / "missing.docx")

    def test_happy_path_paragraphs(self, tmp_path: Path):
        """Lines 64-69: paragraphs are extracted from a real DOCX."""
        docx = _create_minimal_docx(
            tmp_path / "sample.docx",
            paragraphs=["Introduction", "This is the body text."],
        )
        text, meta = parse_docx_text(docx)
        assert "Introduction" in text
        assert "This is the body text." in text
        assert meta["parser"] == "python-docx"
        assert meta["page_count"] == 0

    def test_happy_path_with_table(self, tmp_path: Path):
        """Lines 71-75: tables in DOCX are flattened into pipe-separated text."""
        docx = _create_docx_with_table(
            tmp_path / "table.docx",
            paragraphs=["Table Report"],
            table_rows=[
                ["Name", "Value"],
                ["Metric A", "42"],
            ],
        )
        text, meta = parse_docx_text(docx)
        assert "Table Report" in text
        assert "Name" in text
        assert "Metric A" in text
        assert "42" in text
        assert meta["parser"] == "python-docx"

    def test_empty_paragraphs_skipped(self, tmp_path: Path):
        """Lines 66-69: blank paragraphs are not included."""
        docx = _create_minimal_docx(
            tmp_path / "blanks.docx",
            paragraphs=["First", "", "  ", "Last"],
        )
        text, _ = parse_docx_text(docx)
        lines = [line.strip() for line in text.split("\n\n") if line.strip()]
        assert lines == ["First", "Last"]


# ======================================================================
# parsers.py -- extract_title fallback (no heading)
# ======================================================================


class TestExtractTitle:
    """Tests for title extraction fallback paths."""

    def test_fallback_to_first_nonempty_line(self):
        """When there is no '#' heading, the first substantial line is used."""
        text = "   \nShort\nSome other line"
        # "Short" is len 5 > 3, so it qualifies
        assert extract_title(text, "fallback") == "Short"

    def test_very_short_lines_skipped(self):
        """Lines with 3 or fewer chars are skipped in the fallback path."""
        text = "ab\nReal title here\nMore"
        assert extract_title(text, "fb") == "Real title here"

    def test_heading_with_multiple_hashes(self):
        """## sub-headings are also found by the '#' scan."""
        text = "## Subsection Title\nBody"
        assert extract_title(text, "fb") == "Subsection Title"

    def test_whitespace_collapsed(self):
        """Multiple spaces in a fallback line are collapsed."""
        text = "This   has   extra   spaces"
        assert extract_title(text, "fb") == "This has extra spaces"

    def test_long_title_truncated(self):
        """Titles longer than 120 chars are truncated."""
        text = "# " + "A" * 200
        result = extract_title(text, "fb")
        assert len(result) == 120


# ======================================================================
# chunking.py -- _split_long_text (lines 68-82)
# ======================================================================


class TestSplitLongText:
    """Tests for the internal ``_split_long_text`` function."""

    def test_text_under_limit_returned_as_single_item(self):
        """Line 65-66: short text passes through unchanged."""
        result = _split_long_text("Short text.", max_chars=100)
        assert result == ["Short text."]

    def test_splits_on_sentence_boundaries(self):
        """Lines 68-82: text exceeding max_chars splits at sentence ends."""
        text = "First sentence. Second sentence. Third sentence. Fourth sentence."
        parts = _split_long_text(text, max_chars=40)
        assert len(parts) >= 2
        # All parts should be within or near the limit
        for part in parts:
            assert len(part) <= 80  # generous upper bound
        # Recombined text contains all sentences
        recombined = " ".join(parts)
        assert "First sentence." in recombined
        assert "Fourth sentence." in recombined

    def test_single_long_sentence_no_split_points(self):
        """A single very long sentence with no '.' '!' '?' stays in one chunk."""
        text = "A" * 500
        parts = _split_long_text(text, max_chars=100)
        # No sentence boundaries exist, so the entire text ends up as one chunk
        assert len(parts) == 1
        assert parts[0] == text

    def test_multiple_sentence_endings(self):
        """Handles ! and ? as sentence delimiters."""
        text = "What is this? It is great! Really? Yes!"
        parts = _split_long_text(text, max_chars=25)
        assert len(parts) >= 2
        recombined = " ".join(parts)
        assert "What is this?" in recombined

    def test_empty_string(self):
        """Empty text returns an empty list."""
        result = _split_long_text("", max_chars=100)
        # "" has length 0, which is <= 100, so it returns [""]
        assert result == [""]


# ======================================================================
# chunking.py -- chunk_document edge cases (line 26)
# ======================================================================


class TestChunkDocumentEdgeCases:
    """Cover the no-paragraph fallback (line 26) and empty doc."""

    def test_single_block_no_paragraph_breaks(self):
        """Line 26: text with no double-newline is treated as one paragraph."""
        doc = _make_doc("One block of text with no paragraph breaks at all.")
        chunks = chunk_document(doc)
        assert len(chunks) == 1
        assert chunks[0].text == "One block of text with no paragraph breaks at all."

    def test_empty_document_returns_no_chunks(self):
        """Completely empty text produces no chunks."""
        doc = _make_doc("")
        chunks = chunk_document(doc)
        assert chunks == []

    def test_whitespace_only_returns_no_chunks(self):
        """Text that is only whitespace produces no chunks."""
        doc = _make_doc("   \n\n  \n  ")
        chunks = chunk_document(doc)
        assert chunks == []

    def test_long_paragraph_triggers_split(self):
        """A single paragraph exceeding max_chars is split by _split_long_text."""
        sentences = " ".join(f"Sentence number {i}." for i in range(30))
        doc = _make_doc(sentences)
        chunks = chunk_document(doc, max_chars=100)
        assert len(chunks) > 1
        recombined = " ".join(c.text for c in chunks)
        assert "Sentence number 0." in recombined
        assert "Sentence number 29." in recombined

    def test_overlap_with_short_chunks(self):
        """Overlap larger than previous chunk text works correctly."""
        doc = _make_doc("Hi.\n\nWorld.")
        chunks = chunk_document(doc, overlap_chars=500)
        # First chunk unmodified
        assert chunks[0].text == "Hi."
        # Second chunk prepends entire previous text
        assert chunks[1].text == "Hi.World."

    def test_chunk_documents_empty_list(self):
        """Chunking an empty document list returns an empty chunk list."""
        assert chunk_documents([]) == []


# ======================================================================
# folder.py -- _safe_read_text error path (lines 88-94)
# ======================================================================


class TestSafeReadText:
    """Cover the error handling wrapper for text file reads."""

    @pytest.mark.skipif(
        sys.platform == "win32",
        reason="chmod 000 does not restrict on Windows",
    )
    def test_returns_none_on_permission_error(self, tmp_path: Path, caplog):
        """Lines 88-94: OSError is caught, logged, and returns None."""
        bad_file = tmp_path / "unreadable.md"
        bad_file.write_text("secret", encoding="utf-8")
        os.chmod(bad_file, 0o000)
        try:
            with caplog.at_level(logging.WARNING, logger="cam_rag.documents.folder"):
                result = _safe_read_text(tmp_path, bad_file)
            assert result is None
            assert any("skipping" in r.message for r in caplog.records)
        finally:
            os.chmod(bad_file, 0o644)


# ======================================================================
# folder.py -- _safe_read_parsed error path (lines 101-109)
# ======================================================================


class TestSafeReadParsed:
    """Cover the error handling wrapper for parsed file reads."""

    def test_returns_none_on_file_not_found(self, tmp_path: Path, caplog):
        """Lines 101-109: FileNotFoundError (an OSError) is caught and logged.

        We create and then delete a PDF so the path is stale when
        ``_safe_read_parsed`` tries to open it, triggering a real
        FileNotFoundError from ``parse_pdf_text``.
        """
        pdf = _create_minimal_pdf(tmp_path / "vanished.pdf")
        # Delete the file so parse_pdf_text raises FileNotFoundError (OSError)
        pdf.unlink()
        with caplog.at_level(logging.WARNING, logger="cam_rag.documents.folder"):
            result = _safe_read_parsed(tmp_path, pdf)
        assert result is None
        assert any("skipping" in r.message for r in caplog.records)

    def test_returns_none_on_docx_file_not_found(self, tmp_path: Path, caplog):
        """Same pattern for DOCX: FileNotFoundError path."""
        docx_path = tmp_path / "vanished.docx"
        _create_minimal_docx(docx_path, paragraphs=["temp"])
        docx_path.unlink()
        with caplog.at_level(logging.WARNING, logger="cam_rag.documents.folder"):
            result = _safe_read_parsed(tmp_path, docx_path)
        assert result is None
        assert any("skipping" in r.message for r in caplog.records)


# ======================================================================
# folder.py -- _read_parsed_document PDF and DOCX (lines 133-144)
# ======================================================================


class TestReadParsedDocument:
    """Cover _read_parsed_document for both PDF and DOCX extensions."""

    def test_pdf_produces_corpus_document(self, tmp_path: Path):
        """Lines 133-136, 144-151: PDF -> CorpusDocument with parser metadata."""
        pdf = _create_minimal_pdf(tmp_path / "report.pdf", text="Clinical findings")
        doc = _read_parsed_document(tmp_path, pdf)
        assert isinstance(doc, CorpusDocument)
        assert "Clinical findings" in doc.text
        assert doc.format == "pdf"
        assert doc.source == "report.pdf"
        assert doc.metadata["path"] == "report.pdf"
        assert "parser" in doc.metadata

    def test_docx_produces_corpus_document(self, tmp_path: Path):
        """Lines 137-138, 144-151: DOCX -> CorpusDocument with parser metadata."""
        docx = _create_minimal_docx(
            tmp_path / "notes.docx",
            paragraphs=["Patient notes", "Follow-up in 2 weeks"],
        )
        doc = _read_parsed_document(tmp_path, docx)
        assert isinstance(doc, CorpusDocument)
        assert "Patient notes" in doc.text
        assert "Follow-up in 2 weeks" in doc.text
        assert doc.format == "docx"
        assert doc.metadata["parser"] == "python-docx"

    def test_unsupported_extension_raises(self, tmp_path: Path):
        """Lines 139-142: unsupported extension raises ValueError."""
        rtf = tmp_path / "doc.rtf"
        rtf.write_text("rich text", encoding="utf-8")
        with pytest.raises(ValueError, match="unsupported parsed document extension"):
            _read_parsed_document(tmp_path, rtf)


# ======================================================================
# folder.py -- read_document_folder with PDF and DOCX (lines 44, 55, 63, 72-75)
# ======================================================================


class TestReadDocumentFolderParsedFiles:
    """Cover parsed-file branches in the main ingestion loop."""

    def test_not_a_directory_raises(self, tmp_path: Path):
        """Line 44: NotADirectoryError when path is a file."""
        f = tmp_path / "file.txt"
        f.write_text("not a dir", encoding="utf-8")
        spec = RAGAppSpec(name="test")
        with pytest.raises(NotADirectoryError, match="not a directory"):
            read_document_folder(f, spec)

    def test_skips_non_file_entries(self, tmp_path: Path):
        """Line 55: subdirectories are skipped."""
        sub = tmp_path / "subdir"
        sub.mkdir()
        (tmp_path / "real.md").write_text("# Real\nContent", encoding="utf-8")
        spec = RAGAppSpec(name="test", supported_extensions=(".md",))
        docs = read_document_folder(tmp_path, spec)
        assert len(docs) == 1
        assert docs[0].source == "real.md"

    def test_skips_unsupported_extensions(self, tmp_path: Path):
        """Line 63: files with unsupported extensions are skipped."""
        (tmp_path / "data.csv").write_text("a,b,c", encoding="utf-8")
        (tmp_path / "doc.md").write_text("# Hello\nWorld", encoding="utf-8")
        spec = RAGAppSpec(name="test", supported_extensions=(".md",))
        docs = read_document_folder(tmp_path, spec)
        assert len(docs) == 1

    def test_pdf_ingestion_via_folder(self, tmp_path: Path):
        """Lines 72-75: PDF files are ingested through _safe_read_parsed."""
        _create_minimal_pdf(tmp_path / "report.pdf", text="PDF content here")
        spec = RAGAppSpec(name="test", supported_extensions=(".pdf",))
        docs = read_document_folder(tmp_path, spec)
        assert len(docs) == 1
        assert docs[0].format == "pdf"
        assert "PDF content here" in docs[0].text

    def test_docx_ingestion_via_folder(self, tmp_path: Path):
        """Lines 72-75: DOCX files are ingested through _safe_read_parsed."""
        _create_minimal_docx(
            tmp_path / "notes.docx",
            paragraphs=["Clinical notes content"],
        )
        spec = RAGAppSpec(name="test", supported_extensions=(".docx",))
        docs = read_document_folder(tmp_path, spec)
        assert len(docs) == 1
        assert docs[0].format == "docx"
        assert "Clinical notes content" in docs[0].text

    def test_empty_parsed_file_skipped(self, tmp_path: Path):
        """Lines 73-75: a PDF that produces empty text is not appended."""
        # Create a PDF with an empty page (no text insertion)
        import pymupdf

        doc = pymupdf.open()
        doc.new_page()
        path = tmp_path / "empty.pdf"
        doc.save(str(path))
        doc.close()

        spec = RAGAppSpec(name="test", supported_extensions=(".pdf",))
        docs = read_document_folder(tmp_path, spec)
        assert len(docs) == 0

    def test_mixed_file_types(self, tmp_path: Path):
        """Multiple file types are ingested together."""
        (tmp_path / "readme.md").write_text("# Guide\nInstructions", encoding="utf-8")
        _create_minimal_pdf(tmp_path / "report.pdf", text="PDF report")
        _create_minimal_docx(tmp_path / "notes.docx", paragraphs=["DOCX notes"])
        spec = RAGAppSpec(
            name="test", supported_extensions=(".md", ".pdf", ".docx")
        )
        docs = read_document_folder(tmp_path, spec)
        formats = {d.format for d in docs}
        assert "md" in formats
        assert "pdf" in formats
        assert "docx" in formats


# ======================================================================
# folder.py -- _read_json_documents various shapes (lines 200-202, 223, 225)
# ======================================================================


class TestReadJsonDocumentsShapes:
    """Cover JSON record shapes not yet tested."""

    def test_string_records(self, tmp_path: Path):
        """Lines 199-202: JSON array of plain strings."""
        f = tmp_path / "strings.json"
        f.write_text(json.dumps(["Alpha content", "Beta content"]), encoding="utf-8")
        docs = _read_json_documents(tmp_path, f)
        assert len(docs) == 2
        assert docs[0].text == "Alpha content"
        assert docs[1].text == "Beta content"
        # String records get empty title, which falls back to filename stem
        assert docs[0].title == "strings"

    def test_non_string_non_dict_records_skipped(self, tmp_path: Path):
        """Line 223: records that are neither string nor dict are skipped."""
        f = tmp_path / "weird.json"
        f.write_text(json.dumps([42, True, None, {"text": "valid"}]), encoding="utf-8")
        docs = _read_json_documents(tmp_path, f)
        assert len(docs) == 1
        assert docs[0].text == "valid"

    def test_empty_text_records_skipped(self, tmp_path: Path):
        """Line 225: records with empty or whitespace-only text are skipped."""
        f = tmp_path / "empty_text.json"
        f.write_text(
            json.dumps([{"text": ""}, {"text": "   "}, {"text": "good"}]),
            encoding="utf-8",
        )
        docs = _read_json_documents(tmp_path, f)
        assert len(docs) == 1
        assert docs[0].text == "good"

    def test_metadata_not_dict_ignored(self, tmp_path: Path):
        """Lines 212-216: when metadata is not a dict, it is set to empty."""
        f = tmp_path / "bad_meta.json"
        f.write_text(
            json.dumps([{"text": "hello", "metadata": "not a dict"}]),
            encoding="utf-8",
        )
        docs = _read_json_documents(tmp_path, f)
        assert len(docs) == 1
        assert docs[0].metadata == {"path": "bad_meta.json"}

    def test_title_from_metadata(self, tmp_path: Path):
        """Lines 217-221: title falls back to metadata['title'] if record has none."""
        f = tmp_path / "meta_title.json"
        f.write_text(
            json.dumps([
                {"text": "content", "metadata": {"title": "From Meta"}},
            ]),
            encoding="utf-8",
        )
        docs = _read_json_documents(tmp_path, f)
        assert docs[0].title == "From Meta"

    def test_single_object_json(self, tmp_path: Path):
        """Line 195: a single JSON object (not array) is wrapped in a list."""
        f = tmp_path / "single.json"
        f.write_text(json.dumps({"text": "only one"}), encoding="utf-8")
        docs = _read_json_documents(tmp_path, f)
        assert len(docs) == 1
        assert docs[0].text == "only one"

    def test_content_key_variant(self, tmp_path: Path):
        """_first_string picks 'content' when 'page_content' is missing."""
        f = tmp_path / "content_key.json"
        f.write_text(
            json.dumps([{"content": "via content key"}]),
            encoding="utf-8",
        )
        docs = _read_json_documents(tmp_path, f)
        assert docs[0].text == "via content key"

    def test_page_content_key_variant(self, tmp_path: Path):
        """_first_string picks 'page_content' first."""
        f = tmp_path / "pc_key.json"
        f.write_text(
            json.dumps([{"page_content": "via page_content", "content": "other"}]),
            encoding="utf-8",
        )
        docs = _read_json_documents(tmp_path, f)
        assert docs[0].text == "via page_content"

    def test_body_key_variant(self, tmp_path: Path):
        """_first_string picks 'body' when earlier keys are missing."""
        f = tmp_path / "body_key.json"
        f.write_text(
            json.dumps([{"body": "via body key"}]),
            encoding="utf-8",
        )
        docs = _read_json_documents(tmp_path, f)
        assert docs[0].text == "via body key"

    def test_document_key_variant(self, tmp_path: Path):
        """_first_string picks 'document' when earlier keys are missing."""
        f = tmp_path / "doc_key.json"
        f.write_text(
            json.dumps([{"document": "via document key"}]),
            encoding="utf-8",
        )
        docs = _read_json_documents(tmp_path, f)
        assert docs[0].text == "via document key"

    def test_dict_record_no_text_keys_skipped(self, tmp_path: Path):
        """A dict record with none of the known text keys is skipped."""
        f = tmp_path / "no_keys.json"
        f.write_text(
            json.dumps([{"unknown_field": "value"}]),
            encoding="utf-8",
        )
        docs = _read_json_documents(tmp_path, f)
        assert docs == []

    def test_string_records_in_jsonl(self, tmp_path: Path):
        """JSONL with string records (not dicts)."""
        f = tmp_path / "strings.jsonl"
        f.write_text('"line one"\n"line two"\n', encoding="utf-8")
        docs = _read_json_documents(tmp_path, f)
        assert len(docs) == 2
        assert docs[0].text == "line one"
        assert docs[1].text == "line two"


# ======================================================================
# folder.py -- _first_string (line 255)
# ======================================================================


class TestFirstString:
    """Cover the ``_first_string`` helper directly."""

    def test_returns_first_matching_string(self):
        """Line 250-254: finds the first key whose value is a string."""
        record = {"a": 1, "b": "hello", "c": "world"}
        assert _first_string(record, "a", "b", "c") == "hello"

    def test_returns_empty_when_no_keys_match(self):
        """Line 255: returns '' when none of the keys yield a string value."""
        record = {"a": 1, "b": None}
        assert _first_string(record, "a", "b", "c") == ""

    def test_skips_non_string_values(self):
        """Non-string values are ignored even if the key is present."""
        record = {"text": 42, "body": "real body"}
        assert _first_string(record, "text", "body") == "real body"

    def test_empty_record(self):
        """Empty record returns empty string."""
        assert _first_string({}, "a", "b") == ""
